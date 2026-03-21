from utils import get_output_path
import os
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document

try:
    from pdf2image import convert_from_path
    from pdf2image.exceptions import PDFInfoNotInstalledError
except Exception:
    convert_from_path = None
    PDFInfoNotInstalledError = None

def _poppler_missing_message():
    return (
        "Poppler is required for PDF to image conversion.\n"
        "Windows setup:\n"
        "1) Download a Windows Poppler build.\n"
        "2) Extract it (e.g., C:\\poppler).\n"
        "3) Add C:\\poppler\\Library\\bin to your PATH.\n"
        "4) Restart the app.\n"
        "Then try the conversion again."
    )

def convert_document(file_path, output_format):

    if file_path.endswith(".csv") and output_format == "xlsx":
        output = get_output_path(file_path, "xlsx")
        df = pd.read_csv(file_path)
        df.to_excel(output, index=False)

    elif file_path.endswith(".txt") and output_format == "pdf":
        output = get_output_path(file_path, "pdf")

        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter

        c = canvas.Canvas(output, pagesize=letter)

        with open(file_path, "r") as f:
            lines = f.readlines()

        y = 750
        for line in lines:
            c.drawString(50, y, line.strip())
            y -= 15

        c.save()

    elif file_path.endswith(".pdf") and output_format == "txt":
        output = get_output_path(file_path, "txt")
        reader = PdfReader(file_path)
        with open(output, "w", encoding="utf-8") as f:
            for page in reader.pages:
                text = page.extract_text() or ""
                f.write(text)
                f.write("\n")

    elif file_path.endswith(".pdf") and output_format == "docx":
        output = get_output_path(file_path, "docx")
        reader = PdfReader(file_path)
        doc = Document()
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                doc.add_paragraph(text)
            else:
                doc.add_paragraph("")
        doc.save(output)

    elif file_path.endswith(".pdf") and output_format in ("png", "jpg"):
        if convert_from_path is None:
            raise RuntimeError(_poppler_missing_message())
        try:
            images = convert_from_path(file_path)
        except Exception as e:
            if PDFInfoNotInstalledError and isinstance(e, PDFInfoNotInstalledError):
                raise RuntimeError(_poppler_missing_message())
            raise
        base = os.path.splitext(file_path)[0]
        for idx, img in enumerate(images, start=1):
            out_path = f"{base}_page{idx}.{output_format}"
            if output_format == "jpg":
                img.convert("RGB").save(out_path, "JPEG")
            else:
                img.save(out_path, "PNG")

    else:
        print("Unsupported conversion")
        return

    print(f"Converted to {output}")
